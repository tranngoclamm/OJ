import json
import os
import hashlib
import base64
from cryptography.fernet import Fernet
from calendar import Calendar, SUNDAY
from collections import defaultdict, namedtuple
from datetime import date, datetime, time, timedelta
from functools import partial
from operator import attrgetter, itemgetter

from django import forms
from django.conf import settings
from django.contrib.auth.context_processors import PermWrapper
from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin
from django.core.cache import cache
from django.core.exceptions import ImproperlyConfigured, ObjectDoesNotExist, PermissionDenied
from django.db import IntegrityError
from django.db.models import BooleanField, Case, Count, F, FloatField, IntegerField, Max, Min, Q, Sum, Value, When
from django.db.models.expressions import CombinedExpression
from django.db.models.query import Prefetch
from django.http import Http404, HttpResponse, HttpResponseForbidden, HttpResponseRedirect
from django.shortcuts import get_object_or_404, redirect, render
from django.template.defaultfilters import date as date_filter, floatformat
from django.template.loader import get_template
from django.urls import reverse
from django.utils import timezone
from django.utils.functional import cached_property
from django.utils.html import escape, format_html
from django.utils.safestring import mark_safe
from django.utils.timezone import make_aware
from django.utils.translation import gettext as _, gettext_lazy
from django.views.generic import FormView, ListView, TemplateView, View
from django.views.generic.detail import DetailView, SingleObjectMixin
from django.views.generic.edit import CreateView, UpdateView
from django.views.generic.list import BaseListView
from icalendar import Calendar as ICalendar, Event
from reversion import revisions
from urllib3 import request

from judge.comments import CommentedDetailView
from judge.contest_format import ICPCContestFormat
from judge.forms import ContestAnnouncementForm, ContestCloneForm, ContestDownloadDataForm, ContestForm, \
    ProposeContestProblemFormSet
from judge.models import Contest, ContestAnnouncement, ContestMoss, ContestParticipation, ContestProblem, ContestTag, \
    Language, Organization, Problem, ProblemClarification, Profile, Submission, Device, ContestSeat
from judge.tasks import on_new_contest, prepare_contest_data, run_moss
from judge.utils.celery import redirect_to_task_status, task_status_by_id, task_status_url_by_id
from judge.utils.cms import parse_csv_ranking
from judge.utils.opengraph import generate_opengraph
from judge.utils.problems import _get_result_data, user_attempted_ids, user_completed_ids
from judge.utils.ranker import ranker
from judge.utils.stats import get_bar_chart, get_pie_chart, get_stacked_bar_chart
from judge.utils.views import DiggPaginatorMixin, QueryStringSortMixin, SingleObjectFormView, TitleMixin, \
    add_file_response, generic_message
from judge.views.register import RegistrationForm
from judge.views.seb import SEBRequiredMixin
from django.contrib import messages
from django.db.models import Exists, OuterRef

__all__ = ['ContestList', 'ContestDetail', 'ContestRanking', 'ContestJoin', 'ContestLeave', 'ContestCalendar',
           'ContestClone', 'ContestStats', 'ContestMossView', 'ContestMossDelete',
           'ContestParticipationList', 'ContestParticipationDisqualify', 'get_contest_ranking_list',
           'base_contest_ranking_list']


def _find_contest(request, key, private_check=True):
    try:
        contest = Contest.objects.get(key=key)
        if private_check and not contest.is_accessible_by(request.user):
            raise ObjectDoesNotExist()
    except ObjectDoesNotExist:
        return generic_message(request, _('No such contest'),
                               _('Could not find a contest with the key "%s".') % key, status=404), False
    return contest, True


class ContestListMixin(object):
    hide_private_contests = False

    def get_queryset(self):
        if self.hide_private_contests is not None:
            if 'hide_private_contests' in self.request.GET:
                self.hide_private_contests = self.request.session['hide_private_contests'] \
                                           = self.request.GET.get('hide_private_contests').lower() == 'true'
            else:
                self.hide_private_contests = self.request.session.get('hide_private_contests', False)

        queryset = Contest.get_visible_contests(self.request.user)
        if self.hide_private_contests:
            queryset = queryset.filter(is_organization_private=False)

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['hide_private_contests'] = self.hide_private_contests
        return context


class ContestList(QueryStringSortMixin, DiggPaginatorMixin, TitleMixin, ContestListMixin, ListView):
    model = Contest
    paginate_by = 20
    template_name = 'contest/list.html'
    title = gettext_lazy('Contests')
    context_object_name = 'past_contests'
    all_sorts = frozenset(('name', 'user_count', 'start_time'))
    default_desc = frozenset(('name', 'user_count'))
    default_sort = '-start_time'

    @cached_property
    def _now(self):
        return timezone.now()

    def _get_queryset(self):
        return super().get_queryset().prefetch_related('tags', 'organizations', 'authors', 'curators', 'testers').annotate(
            has_rooms=Exists(
                Contest.exam_room.through.objects.filter(contest_id=OuterRef('pk'))
            )
        )

    def get_queryset(self):
        self.search_query = None
        query_set = self._get_queryset().order_by(self.order, 'key').filter(end_time__lt=self._now)
        if 'search' in self.request.GET:
            self.search_query = search_query = ' '.join(self.request.GET.getlist('search')).strip()
            if search_query:
                query_set = query_set.filter(Q(key__icontains=search_query) | Q(name__icontains=search_query))
        return query_set

    def get_paginator(self, queryset, per_page, orphans=0, allow_empty_first_page=True, **kwargs):
        return super().get_paginator(queryset, per_page, orphans, allow_empty_first_page,
                                     count=self.get_queryset().values('id').count(), **kwargs)

    def get_context_data(self, **kwargs):
        context = super(ContestList, self).get_context_data(**kwargs)
        present, active, future = [], [], []
        finished = set()
        for contest in self._get_queryset().exclude(end_time__lt=self._now):
            if contest.start_time > self._now:
                future.append(contest)
            else:
                present.append(contest)

        if self.request.user.is_authenticated:
            for participation in ContestParticipation.objects.filter(virtual=0, user=self.request.profile,
                                                                     contest_id__in=present) \
                    .select_related('contest') \
                    .prefetch_related('contest__authors', 'contest__curators', 'contest__testers') \
                    .annotate(key=F('contest__key')):
                if participation.ended:
                    finished.add(participation.contest.key)
                else:
                    active.append(participation)
                    present.remove(participation.contest)

        active.sort(key=attrgetter('end_time', 'key'))
        present.sort(key=attrgetter('end_time', 'key'))
        future.sort(key=attrgetter('start_time'))
        context['active_participations'] = active
        context['current_contests'] = present
        context['future_contests'] = future
        context['finished_contests'] = finished
        context['now'] = self._now
        context['first_page_href'] = '.'
        context['page_suffix'] = '#past-contests'
        context['search_query'] = self.search_query
        context.update(self.get_sort_context())
        context.update(self.get_sort_paginate_context())
        return context


class PrivateContestError(Exception):
    def __init__(self, name, is_private, is_organization_private, orgs):
        self.name = name
        self.is_private = is_private
        self.is_organization_private = is_organization_private
        self.orgs = orgs


class ContestMixin(object):
    context_object_name = 'contest'
    model = Contest
    slug_field = 'key'
    slug_url_kwarg = 'contest'

    @cached_property
    def is_in_contest(self):
        return self.object.is_in_contest(self.request.user)

    @cached_property
    def is_editor(self):
        if not self.request.user.is_authenticated:
            return False
        return self.request.profile.id in self.object.editor_ids

    @cached_property
    def is_tester(self):
        if not self.request.user.is_authenticated:
            return False
        return self.request.profile.id in self.object.tester_ids

    @cached_property
    def can_edit(self):
        return self.object.is_editable_by(self.request.user)

    @cached_property
    def can_view_all_problems(self):
        return self.is_in_contest or self.is_editor or self.is_tester or self.request.user.is_superuser or \
            not Problem.objects.filter(contests__contest=self.object, is_public=False).exists()

    def get_context_data(self, **kwargs):
        context = super(ContestMixin, self).get_context_data(**kwargs)
        if self.request.user.is_authenticated:
            try:
                context['live_participation'] = (
                    self.request.profile.contest_history.get(
                        contest=self.object,
                        virtual=ContestParticipation.LIVE,
                    )
                )
            except ContestParticipation.DoesNotExist:
                context['live_participation'] = None
                context['has_joined'] = False
            else:
                context['has_joined'] = True
        else:
            context['live_participation'] = None
            context['has_joined'] = False

        context['now'] = self.object._now
        context['is_in_contest'] = self.is_in_contest
        context['is_editor'] = self.is_editor
        context['is_tester'] = self.is_tester
        context['can_edit'] = self.can_edit

        if not self.object.og_image or not self.object.summary:
            metadata = generate_opengraph('generated-meta-contest:%d' % self.object.id,
                                          self.object.description, 'contest')
        context['meta_description'] = self.object.summary or metadata[0]
        context['og_image'] = self.object.og_image or metadata[1]
        context['has_moss_api_key'] = settings.MOSS_API_KEY is not None
        context['logo_override_image'] = self.object.logo_override_image
        if not context['logo_override_image'] and self.object.organizations.count() == 1:
            context['logo_override_image'] = self.object.organizations.first().logo_override_image

        context['is_ICPC_format'] = (self.object.format.name == ICPCContestFormat.name)
        context['contest_has_rooms'] = self.object.exam_room.filter(code__isnull=False).exists()

        return context

    def get_object(self, queryset=None):
        contest = super(ContestMixin, self).get_object(queryset)

        profile = self.request.profile
        if (profile is not None and
                ContestParticipation.objects.filter(id=profile.current_contest_id, contest_id=contest.id).exists()):
            return contest

        try:
            contest.access_check(self.request.user)
        except Contest.PrivateContest:
            raise PrivateContestError(contest.name, contest.is_private, contest.is_organization_private,
                                      contest.organizations.all())
        except Contest.Inaccessible:
            raise Http404()
        else:
            return contest

    def dispatch(self, request, *args, **kwargs):
        try:
            return super(ContestMixin, self).dispatch(request, *args, **kwargs)
        except Http404:
            key = kwargs.get(self.slug_url_kwarg, None)
            if key:
                return generic_message(request, _('No such contest'),
                                       _('Could not find a contest with the key "%s".') % key)
            else:
                return generic_message(request, _('No such contest'),
                                       _('Could not find such contest.'))
        except PrivateContestError as e:
            return render(request, 'contest/private.html', {
                'error': e, 'title': _('Access to contest "%s" denied') % e.name,
            }, status=403)
        except PermissionDenied as e:
            return generic_message(request, _('Permission denied'), e)

from django.conf import settings
from django.shortcuts import render

class ContestDetail(SEBRequiredMixin, ContestMixin, TitleMixin, CommentedDetailView):
    template_name = 'contest/contest.html'

    def is_comment_locked(self):
        if self.object.use_clarifications:
            now = timezone.now()
            if self.is_in_contest or (self.object.start_time <= now and now <= self.object.end_time):
                return True

        return super(ContestDetail, self).is_comment_locked()

    def get_comment_page(self):
        return 'c:%s' % self.object.key

    def get_title(self):
        return self.object.name

    def get_context_data(self, **kwargs):
        context = super(ContestDetail, self).get_context_data(**kwargs)
        context['can_view_all_problems'] = self.can_view_all_problems
        context['contest_problems'] = Problem.objects.filter(contests__contest=self.object) \
            .order_by('contests__order').defer('description') \
            .annotate(has_public_editorial=Case(
                When(solution__is_public=True, solution__publish_on__lte=timezone.now(), then=True),
                default=False,
                output_field=BooleanField(),
            )) \
            .add_i18n_name(self.request.LANGUAGE_CODE)

        # convert to problem points in contest instead of actual points
        points_list = list(self.object.contest_problems.values_list('points').order_by('order'))
        for idx, p in enumerate(context['contest_problems']):
            p.points = points_list[idx][0]

        context['metadata'] = {
            'has_public_editorials': any(
                problem.is_public and problem.has_public_editorial for problem in context['contest_problems']
            ) if self.object.ended else False,
        }
        context['metadata'].update(
            **self.object.contest_problems
            .annotate(
                partials_enabled=F('partial').bitand(F('problem__partial')),
                pretests_enabled=F('is_pretested').bitand(F('contest__run_pretests_only')),
            )
            .aggregate(
                has_partials=Sum('partials_enabled', output_field=BooleanField()),
                has_pretests=Sum('pretests_enabled', output_field=BooleanField()),
                has_submission_cap=Sum('max_submissions'),
                problem_count=Count('id'),
            ),
        )

        clarifications = ProblemClarification.objects.filter(problem__in=self.object.problems.all())
        context['has_clarifications'] = clarifications.count() > 0
        context['clarifications'] = clarifications.order_by('-date')
        announcements = ContestAnnouncement.objects.filter(contest=self.object)
        context['has_announcements'] = announcements.count() > 0
        context['announcements'] = announcements.order_by('-date')
        context['can_announce'] = self.object.is_editable_by(self.request.user)

        authenticated = self.request.user.is_authenticated
        context['completed_problem_ids'] = user_completed_ids(self.request.profile) if authenticated else []
        context['attempted_problem_ids'] = user_attempted_ids(self.request.profile) if authenticated else []

        context['can_download_data'] = bool(settings.DMOJ_CONTEST_DATA_DOWNLOAD)

        return context
    
    def dispatch(self, request, *args, **kwargs):
        # Chỉ áp dụng nếu user đang đăng nhập và contest là kỳ thi đặc biệt
        user = request.user
        if user.is_authenticated:
            contest = self.get_object()
            response = self.seb_check(request, contest)
            if response:
                return response
        return super().dispatch(request, *args, **kwargs)


class ContestAllProblems(ContestMixin, TitleMixin, DetailView):
    template_name = 'contest/contest-all-problems.html'

    def get_title(self):
        return self.object.name

    def get_context_data(self, **kwargs):
        context = super(ContestAllProblems, self).get_context_data(**kwargs)

        if not self.can_view_all_problems:
            raise Http404()

        context['contest_problems'] = Problem.objects.filter(contests__contest=self.object) \
            .order_by('contests__order') \
            .add_i18n_name(self.request.LANGUAGE_CODE) \
            .add_i18n_description(self.request.LANGUAGE_CODE)

        # convert to problem points in contest instead of actual points
        points_list = list(self.object.contest_problems.values_list('points').order_by('order'))
        for idx, p in enumerate(context['contest_problems']):
            p.points = points_list[idx][0]

        authenticated = self.request.user.is_authenticated
        context['completed_problem_ids'] = user_completed_ids(self.request.profile) if authenticated else []
        context['attempted_problem_ids'] = user_attempted_ids(self.request.profile) if authenticated else []

        return context


class ContestClone(ContestMixin, PermissionRequiredMixin, TitleMixin, SingleObjectFormView):
    title = gettext_lazy('Clone Contest')
    template_name = 'contest/clone.html'
    form_class = ContestCloneForm
    permission_required = 'judge.clone_contest'
    permission_denied_message = _('You are not allowed to clone contests.')

    def get_object(self, queryset=None):
        contest = super().get_object(queryset)
        if not contest.is_editable_by(self.request.user):
            raise PermissionDenied(_('You are not allowed to edit this contest.'))
        return contest

    def form_valid(self, form):
        contest = self.object

        # Using list() to force QuerySets evaluation, as `contest.pk = None` affects these queries
        tags = list(contest.tags.all())
        organizations = list(contest.organizations.all())
        private_contestants = list(contest.private_contestants.all())
        view_contest_scoreboard = list(contest.view_contest_scoreboard.all())
        contest_problems = list(contest.contest_problems.all())
        old_key = contest.key

        contest.pk = None
        contest.is_visible = False
        contest.user_count = 0
        contest.virtual_count = 0
        contest.locked_after = None
        contest.key = form.cleaned_data['key']
        with revisions.create_revision(atomic=True):
            contest.save()
            contest.tags.set(tags)
            contest.organizations.set(organizations)
            contest.private_contestants.set(private_contestants)
            contest.view_contest_scoreboard.set(view_contest_scoreboard)
            contest.authors.add(self.request.profile)

            for problem in contest_problems:
                problem.contest = contest
                problem.pk = None
            ContestProblem.objects.bulk_create(contest_problems)

            revisions.set_user(self.request.user)
            revisions.set_comment(_('Cloned contest from %s') % old_key)

        return HttpResponseRedirect(reverse('contest_edit', args=(contest.key,)))


class ContestAnnounce(ContestMixin, TitleMixin, SingleObjectFormView):
    title = gettext_lazy('Create contest announcement')
    template_name = 'contest/create-announcement.html'
    form_class = ContestAnnouncementForm

    def get_object(self, queryset=None):
        contest = super().get_object(queryset)
        if not contest.is_editable_by(self.request.user):
            raise PermissionDenied(_('You are not allowed to edit this contest.'))
        return contest

    def form_valid(self, form):
        contest = self.object

        announcement = form.save(commit=False)
        announcement.contest = contest
        announcement.save()
        announcement.send()

        return HttpResponseRedirect(reverse('contest_view', args=(contest.key,)))


class ContestAccessDenied(Exception):
    pass


class ContestAccessCodeForm(forms.Form):
    access_code = forms.CharField(max_length=255)

    def __init__(self, *args, **kwargs):
        super(ContestAccessCodeForm, self).__init__(*args, **kwargs)
        self.fields['access_code'].widget.attrs.update({'autocomplete': 'off'})


class ContestRegister(LoginRequiredMixin, ContestMixin, SingleObjectMixin, View):
    def get(self, request, *args, **kwargs):
        self.object = self.get_object()
        return self.ask_for_access_code()

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        try:
            return self.register_contest(request)
        except ContestAccessDenied:
            if request.POST.get('access_code'):
                return self.ask_for_access_code(ContestAccessCodeForm(request.POST))
            else:
                return HttpResponseRedirect(request.path)

    def register_contest(self, request, access_code=None):
        contest = self.object
        profile = request.profile

        if self.is_editor or self.is_tester:
            return generic_message(request, _('Cannot register'),
                                   _('You cannot register for this contest.'))

        if not request.user.is_superuser and contest.banned_users.filter(id=profile.id).exists():
            return generic_message(request, _('Banned from joining'),
                                   _('You have been declared persona non grata for this contest. '
                                     'You are permanently barred from joining this contest.'))

        if not contest.require_registration:
            return generic_message(request, _('Cannot register'),
                                   _('Registration is not required for this contest.'))

        if not contest.can_register:
            return generic_message(request, _('Cannot register'),
                                   _('You cannot register for this contest now.'))

        requires_access_code = (not self.can_edit and contest.access_code and access_code != contest.access_code)
        if contest.ended:
            return generic_message(request, _('Contest has ended'),
                                   _('"%s" has ended.') % contest.name)
        else:
            if self.is_editor or self.is_tester:
                return generic_message(request, _('Cannot register'),
                                       _('You cannot register for this contest.'))

            try:
                ContestParticipation.objects.get(
                    contest=contest, user=profile, virtual=0,
                )
            except ContestParticipation.DoesNotExist:
                if requires_access_code:
                    raise ContestAccessDenied()

                ContestParticipation.objects.create(
                    contest=contest, user=profile, virtual=0,
                    real_start=datetime(1970, 1, 1, tzinfo=timezone.utc),
                )
            else:
                return generic_message(request, _('Already registered'),
                                       _('You have already registered for this contest.'))

        contest._updating_stats_only = True
        contest.update_user_count()
        return HttpResponseRedirect(reverse('contest_view', args=(contest.key,)))

    def ask_for_access_code(self, form=None):
        contest = self.object
        wrong_code = False
        if form:
            if form.is_valid():
                if form.cleaned_data['access_code'] == contest.access_code:
                    return self.register_contest(self.request, form.cleaned_data['access_code'])
                wrong_code = True
        else:
            form = ContestAccessCodeForm()
        return render(self.request, 'contest/access_code.html', {
            'form': form, 'wrong_code': wrong_code,
            'title': _('Enter access code for "%s"') % contest.name,
        })


class ContestJoin(LoginRequiredMixin, SEBRequiredMixin, ContestMixin, SingleObjectMixin, View):
    def get(self, request, *args, **kwargs):
        self.object = self.get_object()
        return self.ask_for_access_code()

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        try:
            return self.join_contest(request)
        except ContestAccessDenied:
            if request.POST.get('access_code'):
                return self.ask_for_access_code(ContestAccessCodeForm(request.POST))
            else:
                return HttpResponseRedirect(request.path)

    def join_contest(self, request, access_code=None):
        contest = self.object

        if not contest.can_join and not (self.is_editor or self.is_tester):
            return generic_message(request, _('Contest not ongoing'),
                                   _('"%s" is not currently ongoing.') % contest.name)

        profile = request.profile

        if not request.user.is_superuser and contest.banned_users.filter(id=profile.id).exists():
            return generic_message(request, _('Banned from joining'),
                                   _('You have been declared persona non grata for this contest. '
                                     'You are permanently barred from joining this contest.'))

        is_privileged = request.user.is_staff or request.user.is_superuser or self.is_editor or self.is_tester
        if request.is_seb and request.in_exam_contest and not request.profile.current_contest.finished_at and not is_privileged:
            return generic_message(request, _('Contest join failed.'),
                                   _('You must end the exam in "%s" before joining this contest.') % request.participation.contest.name)

        # Conditions for joining a contest:
        #   - If contest has ended, allow virtual joining iff:
        #       - contest.disallow_virtual is False
        #       - requires_access_code is False
        #   - If contest is ongoing, allow joining iff:
        #       - Not editor or tester
        #       - Registered if registration windows has ended
        #       - requires_access_code is False
        #   - Editors/Testers can only spectate live contests and only when requires_access_code is False.

        requires_access_code = (not self.can_edit and contest.access_code and access_code != contest.access_code)
        if contest.ended:
            if contest.disallow_virtual:
                return generic_message(request, _('Virtual joining not allowed'),
                                       _('Virtual joining is not allowed for this contest.'))

            if requires_access_code:
                raise ContestAccessDenied()

            while True:
                virtual_id = max((ContestParticipation.objects.filter(contest=contest, user=profile)
                                  .aggregate(virtual_id=Max('virtual'))['virtual_id'] or 0) + 1, 1)
                try:
                    participation = ContestParticipation.objects.create(
                        contest=contest, user=profile, virtual=virtual_id,
                        real_start=timezone.now(),
                    )
                # There is obviously a race condition here, so we keep trying until we win the race.
                except IntegrityError:
                    pass
                else:
                    break
        else:
            SPECTATE = ContestParticipation.SPECTATE
            LIVE = ContestParticipation.LIVE
            can_only_spectate = self.is_editor or self.is_tester
            try:
                participation = ContestParticipation.objects.get(
                    contest=contest, user=profile, virtual=(SPECTATE if can_only_spectate else LIVE),
                )
            except ContestParticipation.DoesNotExist:
                if contest.require_registration and not contest.can_register and not can_only_spectate:
                    return generic_message(request, _('Not registered'),
                                           _('You are not registered for this contest.'))

                if requires_access_code:
                    raise ContestAccessDenied()

                participation = ContestParticipation.objects.create(
                    contest=contest, user=profile, virtual=(SPECTATE if can_only_spectate else LIVE),
                    real_start=timezone.now(),
                )
            else:
                if participation.pre_registered:
                    # Pre-registered. First time joining.
                    participation.real_start = timezone.now()
                    participation.save()

                if participation.ended:
                    participation = ContestParticipation.objects.get_or_create(
                        contest=contest, user=profile, virtual=SPECTATE,
                        defaults={'real_start': timezone.now()},
                    )[0]

        profile.current_contest = participation
        profile.save()
        contest._updating_stats_only = True
        contest.update_user_count()
        return HttpResponseRedirect(reverse('contest_view', args=(contest.key,)))

    def ask_for_access_code(self, form=None):
        contest = self.object
        wrong_code = False
        if form:
            if form.is_valid():
                if form.cleaned_data['access_code'] == contest.access_code:
                    return self.join_contest(self.request, form.cleaned_data['access_code'])
                wrong_code = True
        else:
            form = ContestAccessCodeForm()
        return render(self.request, 'contest/access_code.html', {
            'form': form, 'wrong_code': wrong_code,
            'title': _('Enter access code for "%s"') % contest.name,
        })

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        user = request.user

        if user.is_authenticated:
            contest = self.get_object()
            response = self.seb_check(request, contest)
            if response:
                return response
        return super().dispatch(request, *args, **kwargs)

class ContestLeave(LoginRequiredMixin, ContestMixin, SingleObjectMixin, View):
    def dispatch(self, request, *args, **kwargs):
        if request.method != 'POST':
            return HttpResponseForbidden()

        return super(ContestLeave, self).dispatch(request, *args, **kwargs)

    def post(self, request, *args, **kwargs):
        contest = self.get_object()

        if request.is_seb and not (request.user.is_staff or request.user.is_superuser):
            return HttpResponseForbidden("Cannot leave contest while using SEB.")
        
        profile = request.profile
        if profile.current_contest is None or profile.current_contest.contest_id != contest.id:
            return generic_message(request, _('No such contest'),
                                   _('You are not in contest "%s".') % contest.key, 404)

        profile.remove_contest()
        return HttpResponseRedirect(reverse('contest_view', args=(contest.key,)))


ContestDay = namedtuple('ContestDay', 'date is_pad is_today starts ends oneday')


class ContestCalendar(TitleMixin, ContestListMixin, TemplateView):
    firstweekday = SUNDAY
    template_name = 'contest/calendar.html'

    def get(self, request, *args, **kwargs):
        try:
            self.year = int(kwargs['year'])
            self.month = int(kwargs['month'])
        except (KeyError, ValueError):
            raise ImproperlyConfigured('ContestCalendar requires integer year and month')
        self.today = timezone.now().date()
        return self.render()

    def render(self):
        context = self.get_context_data()
        return self.render_to_response(context)

    def get_contest_data(self, start, end):
        end += timedelta(days=1)
        contests = self.get_queryset().filter(Q(start_time__gte=start, start_time__lt=end) |
                                              Q(end_time__gte=start, end_time__lt=end))
        starts, ends, oneday = (defaultdict(list) for i in range(3))
        for contest in contests:
            start_date = timezone.localtime(contest.start_time).date()
            end_date = timezone.localtime(contest.end_time - timedelta(seconds=1)).date()
            if start_date == end_date:
                oneday[start_date].append(contest)
            else:
                starts[start_date].append(contest)
                ends[end_date].append(contest)
        return starts, ends, oneday

    def get_table(self):
        calendar = Calendar(self.firstweekday).monthdatescalendar(self.year, self.month)
        starts, ends, oneday = self.get_contest_data(make_aware(datetime.combine(calendar[0][0], time.min)),
                                                     make_aware(datetime.combine(calendar[-1][-1], time.min)))
        return [[ContestDay(
            date=date, is_pad=date.month != self.month,
            is_today=date == self.today, starts=starts[date], ends=ends[date], oneday=oneday[date],
        ) for date in week] for week in calendar]

    def get_context_data(self, **kwargs):
        context = super(ContestCalendar, self).get_context_data(**kwargs)

        try:
            month = date(self.year, self.month, 1)
        except ValueError:
            raise Http404()
        else:
            context['title'] = _('Contests in %(month)s') % {'month': date_filter(month, _('F Y'))}

        dates = Contest.objects.aggregate(min=Min('start_time'), max=Max('end_time'))
        min_month = (self.today.year, self.today.month)
        if dates['min'] is not None:
            min_month = dates['min'].year, dates['min'].month
        max_month = (self.today.year, self.today.month)
        if dates['max'] is not None:
            max_month = max((dates['max'].year, dates['max'].month), (self.today.year, self.today.month))

        month = (self.year, self.month)
        if month < min_month or month > max_month:
            # 404 is valid because it merely declares the lack of existence, without any reason
            raise Http404()

        context['now'] = timezone.now()
        context['calendar'] = self.get_table()
        context['curr_month'] = date(self.year, self.month, 1)

        if month > min_month:
            context['prev_month'] = date(self.year - (self.month == 1), 12 if self.month == 1 else self.month - 1, 1)
        else:
            context['prev_month'] = None

        if month < max_month:
            context['next_month'] = date(self.year + (self.month == 12), 1 if self.month == 12 else self.month + 1, 1)
        else:
            context['next_month'] = None
        return context


class ContestICal(TitleMixin, ContestListMixin, BaseListView):
    def generate_ical(self):
        cal = ICalendar()
        cal.add('prodid', '-//DMOJ//NONSGML Contests Calendar//')
        cal.add('version', '2.0')

        now = timezone.now().astimezone(timezone.utc)
        domain = self.request.get_host()
        for contest in self.get_queryset():
            event = Event()
            event.add('uid', f'contest-{contest.key}@{domain}')
            event.add('summary', contest.name)
            event.add('location', self.request.build_absolute_uri(contest.get_absolute_url()))
            event.add('dtstart', contest.start_time.astimezone(timezone.utc))
            event.add('dtend', contest.end_time.astimezone(timezone.utc))
            event.add('dtstamp', now)
            cal.add_component(event)
        return cal.to_ical()

    def render_to_response(self, context, **kwargs):
        return HttpResponse(self.generate_ical(), content_type='text/calendar')


class ContestStats(TitleMixin, ContestMixin, DetailView):
    template_name = 'contest/stats.html'

    def get_title(self):
        return _('%s Statistics') % self.object.name

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        if not self.object.can_see_full_submission_list(self.request.user):
            raise Http404()

        queryset = Submission.objects.filter(contest_object=self.object, date__gt=self.object.start_time)

        ac_count = Count(Case(When(result='AC', then=Value(1)), output_field=IntegerField()))
        ac_rate = CombinedExpression(ac_count / Count('problem'), '*', Value(100.0), output_field=FloatField())

        status_count_queryset = list(
            queryset.values('problem__code', 'result').annotate(count=Count('result'))
                    .values_list('problem__code', 'result', 'count'),
        )
        labels, codes = [], []
        contest_problems = self.object.contest_problems.order_by('order').values_list('problem__name', 'problem__code')
        if contest_problems:
            labels, codes = zip(*contest_problems)
        num_problems = len(labels)
        status_counts = [[] for i in range(num_problems)]
        for problem_code, result, count in status_count_queryset:
            if problem_code in codes:
                status_counts[codes.index(problem_code)].append((result, count))

        result_data = defaultdict(partial(list, [0] * num_problems))
        for i in range(num_problems):
            for category in _get_result_data(defaultdict(int, status_counts[i]))['categories']:
                result_data[category['code']][i] = category['count']

        language_id_to_name = {id: name for id, name in Language.objects.values_list('id', 'name')}

        def id_to_name(data):
            return (language_id_to_name[data[0]], data[1])

        stats = {
            'problem_status_count': get_stacked_bar_chart(
                labels, result_data, settings.DMOJ_STATS_SUBMISSION_RESULT_COLORS,
            ),
            'problem_ac_rate': get_bar_chart(
                queryset.values('contest__problem__order', 'problem__name').annotate(ac_rate=ac_rate)
                        .order_by('contest__problem__order').values_list('problem__name', 'ac_rate'),
            ),
            'language_count': get_pie_chart(
                list(map(id_to_name, queryset.values('language_id').annotate(count=Count('language_id'))
                         .filter(count__gt=0).order_by('-count').values_list('language_id', 'count'))),
            ),
            'language_ac_rate': get_bar_chart(
                list(map(id_to_name, queryset.values('language_id').annotate(ac_rate=ac_rate)
                         .filter(ac_rate__gt=0).values_list('language_id', 'ac_rate'))),
            ),
        }

        context['stats'] = mark_safe(json.dumps(stats))

        return context


ContestRankingProfile = namedtuple(
    'ContestRankingProfile',
    'id user css_class username points cumtime tiebreaker organization participation '
    'participation_rating problem_cells result_cell virtual display_name',
)

BestSolutionData = namedtuple('BestSolutionData', 'code points time state is_pretested')


def make_contest_ranking_profile(contest, participation, contest_problems, first_solves, frozen=False):
    def display_user_problem(contest_problem):
        # When the contest format is changed, `format_data` might be invalid.
        # This will cause `display_user_problem` to error, so we display '???' instead.
        try:
            return contest.format.display_user_problem(participation, contest_problem, first_solves, frozen)
        except (KeyError, TypeError, ValueError):
            return mark_safe('<td>???</td>')

    user = participation.user
    return ContestRankingProfile(
        id=user.id,
        user=user.user,
        css_class=user.css_class,
        username=user.username,
        points=participation.score if not frozen else participation.frozen_score,
        cumtime=participation.cumtime if not frozen else participation.frozen_cumtime,
        tiebreaker=participation.tiebreaker if not frozen else participation.frozen_tiebreaker,
        organization=user.organization,
        participation_rating=participation.rating.rating if hasattr(participation, 'rating') else None,
        problem_cells=[display_user_problem(contest_problem) for contest_problem in contest_problems],
        result_cell=contest.format.display_participation_result(participation, frozen),
        participation=participation,
        virtual=participation.virtual,
        display_name=user.display_name,
    )


def base_contest_ranking_list(contest, problems, queryset, frozen=False):
    queryset = queryset.select_related('user__user', 'rating').defer('user__about', 'user__organizations__about')
    first_solves, total_ac = contest.format.get_first_solves_and_total_ac(problems, queryset, frozen)
    users = [make_contest_ranking_profile(contest, participation, problems, first_solves, frozen) for participation
             in queryset]
    return users, total_ac


def base_contest_ranking_queryset(contest):
    return contest.users.filter(virtual__gt=ContestParticipation.SPECTATE) \
        .prefetch_related(Prefetch('user__organizations',
                                   queryset=Organization.objects.filter(is_unlisted=False))) \
        .annotate(submission_count=Count('submission')) \
        .order_by('is_disqualified', '-score', 'cumtime', 'tiebreaker', '-submission_count')


def base_contest_frozen_ranking_queryset(contest):
    return contest.users.filter(virtual__gt=ContestParticipation.SPECTATE) \
        .prefetch_related(Prefetch('user__organizations',
                                   queryset=Organization.objects.filter(is_unlisted=False))) \
        .annotate(submission_count=Count('submission')) \
        .order_by('is_disqualified', '-frozen_score', 'frozen_cumtime', 'frozen_tiebreaker', '-submission_count')


def contest_ranking_list(contest, problems, frozen=False):
    return base_contest_ranking_list(contest, problems, base_contest_ranking_queryset(contest), frozen=frozen)


def get_contest_ranking_list(request, contest, participation=None, ranking_list=contest_ranking_list, ranker=ranker):
    problems = list(contest.contest_problems.select_related('problem').defer('problem__description').order_by('order'))
    users, total_ac = ranking_list(contest, problems)
    users = ranker(users, key=attrgetter('points', 'cumtime', 'tiebreaker'))

    return users, problems, total_ac


class ContestRankingBase(ContestMixin, TitleMixin, DetailView):
    template_name = 'contest/ranking.html'
    ranking_table_template = get_template('contest/ranking-table.html')
    tab = None

    def get_title(self):
        raise NotImplementedError()

    def get_content_title(self):
        return self.object.name

    def get_ranking_list(self):
        raise NotImplementedError()

    @property
    def is_frozen(self):
        return False

    def check_can_see_own_scoreboard(self):
        if not self.object.can_see_own_scoreboard(self.request.user):
            raise Http404()

    def get_rendered_ranking_table(self):
        users, problems, total_ac = self.get_ranking_list()

        return self.ranking_table_template.render(request=self.request, context={
            'table_id': 'ranking-table',
            'users': users,
            'problems': problems,
            'total_ac': total_ac,
            'contest': self.object,
            'has_rating': self.object.ratings.exists(),
            'is_frozen': self.is_frozen,
            'perms': PermWrapper(self.request.user),
            'can_edit': self.can_edit,
            'is_ICPC_format': (self.object.format.name == ICPCContestFormat.name),
        })

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        self.check_can_see_own_scoreboard()

        context['rendered_ranking_table'] = self.get_rendered_ranking_table()
        context['tab'] = self.tab
        return context

    def get(self, request, *args, **kwargs):
        if 'raw' in request.GET:
            self.object = self.get_object()

            self.check_can_see_own_scoreboard()

            return HttpResponse(self.get_rendered_ranking_table(), content_type='text/plain')

        return super().get(request, *args, **kwargs)


class ContestRanking(ContestRankingBase):
    tab = 'ranking'
    show_virtual = False

    def get_title(self):
        return _('%s Rankings') % self.object.name

    @cached_property
    def is_frozen(self):
        return self.object.is_frozen and not self.can_edit

    @property
    def cache_key(self):
        return f'contest_ranking_cache_{self.object.key}_{self.show_virtual}_{self.is_frozen}_' \
               f'{self.request.LANGUAGE_CODE}'

    @property
    def bypass_cache_ranking(self):
        return self.object.scoreboard_cache_timeout == 0 or self.can_edit or \
            (self.request.user.is_authenticated and not self.object.can_see_full_scoreboard(self.request.user))

    def get_ranking_queryset(self):
        if self.is_frozen:
            queryset = base_contest_frozen_ranking_queryset(self.object)
        else:
            queryset = base_contest_ranking_queryset(self.object)
        if not self.show_virtual:
            queryset = queryset.filter(virtual=ContestParticipation.LIVE)
        return queryset

    def get_full_ranking_list(self):
        if 'show_virtual' in self.request.GET:
            self.show_virtual = self.request.session['show_virtual'] \
                              = self.request.GET.get('show_virtual').lower() == 'true'
        else:
            self.show_virtual = self.request.session.get('show_virtual', False)

        queryset = self.get_ranking_queryset()
        return get_contest_ranking_list(
            self.request, self.object,
            ranking_list=partial(base_contest_ranking_list, queryset=queryset, frozen=self.is_frozen),
        )

    def get_ranking_list(self):
        if not self.object.can_see_full_scoreboard(self.request.user):
            queryset = self.object.users.filter(user=self.request.profile, virtual=ContestParticipation.LIVE)
            return get_contest_ranking_list(
                self.request, self.object,
                ranking_list=partial(base_contest_ranking_list, queryset=queryset),
                ranker=lambda users, key: ((_('???'), user) for user in users),
            )

        return self.get_full_ranking_list()

    def get_rendered_ranking_table(self):
        if self.bypass_cache_ranking:
            return super().get_rendered_ranking_table()

        rendered_ranking_table = cache.get(self.cache_key, None)
        if rendered_ranking_table is None:
            rendered_ranking_table = super().get_rendered_ranking_table()
            cache.set(self.cache_key, rendered_ranking_table, self.object.scoreboard_cache_timeout)

        return rendered_ranking_table

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['has_rating'] = self.object.ratings.exists()
        context['show_virtual'] = self.show_virtual
        context['is_frozen'] = self.is_frozen
        context['cache_timeout'] = 0 if self.bypass_cache_ranking else self.object.scoreboard_cache_timeout
        return context


class ContestPublicRanking(ContestRanking):
    def check_can_see_own_scoreboard(self):
        # ignore this check, we want to show the scoreboard to everyone
        pass

    def get_ranking_list(self):
        # ignore the `can_see_full_scoreboard` check
        return self.get_full_ranking_list()

    def get(self, request, *args, **kwargs):
        self.object = self.get_object()
        ranking_access_code = self.object.ranking_access_code
        if not ranking_access_code or ranking_access_code != request.GET.get('code'):
            return generic_message(request, _('Ranking access code required'),
                                   _('You need to provide a valid ranking access code to access this page.'))

        return super().get(request, *args, **kwargs)


class ContestOfficialRanking(ContestRankingBase):
    template_name = 'contest/official-ranking.html'
    ranking_table_template = get_template('contest/official-ranking-table.html')
    tab = 'official_ranking'

    def get_title(self):
        return _('%s Official Rankings') % self.object.name

    def get_ranking_list(self):
        def display_points(points):
            return format_html(
                '<td class="user-points">{points}</td>',
                points=floatformat(points),
            )

        users, problems = parse_csv_ranking(self.object.csv_ranking)

        for user in users:
            user['result_cell'] = display_points(user['total_score'])
            user['problem_cells'] = [display_points(points) for points in user['scores']]

        users = list(zip(range(1, len(users) + 1), users))

        return users, problems, {}

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['has_rating'] = False
        return context

    def get(self, request, *args, **kwargs):
        self.object = self.get_object()
        if not self.object.csv_ranking:
            raise Http404()

        # If the csv_ranking is an url, redirect to it
        # (the check is not perfect, but it's good enough)
        if self.object.csv_ranking.startswith('http'):
            return redirect(self.object.csv_ranking)

        return super().get(request, *args, **kwargs)


class ContestParticipationList(LoginRequiredMixin, ContestRankingBase):
    tab = 'participation'

    def get_title(self):
        if self.profile == self.request.profile:
            return _('Your participation in %(contest)s') % {'contest': self.object.name}
        return _("%(user)s's participation in %(contest)s") % {
            'user': self.profile.username, 'contest': self.object.name,
        }

    def get_ranking_list(self):
        if not self.object.can_see_full_scoreboard(self.request.user) and self.profile != self.request.profile:
            raise Http404()

        queryset = self.object.users.filter(user=self.profile, virtual__gte=0).order_by('-virtual')
        live_link = format_html('<a href="{2}#!{1}">{0}</a>', _('Live'), self.profile.username,
                                reverse('contest_ranking', args=[self.object.key]))

        return get_contest_ranking_list(
            self.request, self.object,
            ranking_list=partial(base_contest_ranking_list, queryset=queryset),
            ranker=lambda users, key: ((user.participation.virtual or live_link, user) for user in users))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['has_rating'] = False
        context['rank_header'] = _('Participation')
        return context

    def get(self, request, *args, **kwargs):
        if 'user' in kwargs:
            self.profile = get_object_or_404(Profile, user__username=kwargs['user'])
        else:
            self.profile = self.request.profile
        return super().get(request, *args, **kwargs)


class ContestParticipationDisqualify(ContestMixin, SingleObjectMixin, View):
    def get_object(self, queryset=None):
        contest = super().get_object(queryset)
        if not contest.is_editable_by(self.request.user):
            raise Http404()
        return contest

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()

        try:
            participation = self.object.users.get(pk=request.POST.get('participation'))
        except ObjectDoesNotExist:
            pass
        else:
            participation.set_disqualified(not participation.is_disqualified)
        return HttpResponseRedirect(reverse('contest_ranking', args=(self.object.key,)))


class ContestMossMixin(ContestMixin, PermissionRequiredMixin):
    permission_required = 'judge.moss_contest'
    permission_denied_message = _('You are not allowed to run MOSS.')

    def get_object(self, queryset=None):
        contest = super().get_object(queryset)
        if settings.MOSS_API_KEY is None or not contest.is_editable_by(self.request.user):
            raise Http404()
        return contest


class ContestMossView(ContestMossMixin, TitleMixin, DetailView):
    template_name = 'contest/moss.html'

    def get_title(self):
        return _('%s MOSS Results') % self.object.name

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        problems = list(map(attrgetter('problem'), self.object.contest_problems.order_by('order')
                                                              .select_related('problem')))
        languages = list(map(itemgetter(0), ContestMoss.LANG_MAPPING))

        results = ContestMoss.objects.filter(contest=self.object)
        moss_results = defaultdict(list)
        for result in results:
            moss_results[result.problem].append(result)

        for result_list in moss_results.values():
            result_list.sort(key=lambda x: languages.index(x.language))

        context['languages'] = languages
        context['has_results'] = results.exists()
        context['moss_results'] = [(problem, moss_results[problem]) for problem in problems]

        return context

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        status = run_moss.delay(self.object.key)
        return redirect_to_task_status(
            status, message=_('Running MOSS for %s...') % (self.object.name,),
            redirect=reverse('contest_moss', args=(self.object.key,)),
        )


class ContestMossDelete(ContestMossMixin, SingleObjectMixin, View):
    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        ContestMoss.objects.filter(contest=self.object).delete()
        return HttpResponseRedirect(reverse('contest_moss', args=(self.object.key,)))


class ContestTagDetailAjax(DetailView):
    model = ContestTag
    slug_field = slug_url_kwarg = 'name'
    context_object_name = 'tag'
    template_name = 'contest/tag-ajax.html'


class ContestTagDetail(TitleMixin, ContestTagDetailAjax):
    template_name = 'contest/tag.html'

    def get_title(self):
        return _('Contest tag: %s') % self.object.name


from django.contrib.auth import get_user_model
from django.utils.text import slugify
from docx import Document
import random
import string
from django.http import HttpResponseRedirect
def random_upper(length):
    return ''.join(random.choices(string.ascii_uppercase, k=length))

from io import BytesIO
from django.http import HttpResponse
from copy import deepcopy 
# from judge.models.exam_access import ExamAccess
from django.http import FileResponse

import threading

def download_account_docx(request):
    path = request.session.pop('account_docx_path', None)
    if not path or not os.path.exists(path):
        raise Http404("File not found.")
    file_handle = open(path, 'rb')
    response = FileResponse(file_handle, as_attachment=True, filename=os.path.basename(path))

    def delayed_delete(file_path):
        import time
        time.sleep(0.5)  
        try:
            os.remove(file_path)
        except Exception as e:
            print(f"Không thể xoá file {file_path}: {e}")

    threading.Thread(target=delayed_delete, args=(path,)).start()

    return response

def export_accounts_to_docx(account_list, filename_prefix):
    doc = Document()
    doc.add_heading(f'Danh sách tài khoản {filename_prefix}', level=1)    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tài khoản'
    hdr_cells[1].text = 'Mật khẩu'

    for username, password in account_list:
        row_cells = table.add_row().cells
        row_cells[0].text = username
        row_cells[1].text = password

    # Tạo tên file duy nhất
    filename = f'{datetime.now().strftime("%Y%m%d")}_{filename_prefix}.docx'
    file_path = os.path.join(settings.MEDIA_ROOT, 'exports', filename)

    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    doc.save(file_path)

    return file_path

def export_accounts_to_docx_with_devices(device_mapping, filename_prefix):
    doc = Document()
    doc.add_heading(f'Danh sách tài khoản + máy thi {filename_prefix}', level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tài khoản'
    hdr_cells[1].text = 'Máy (hostname)'

    for username_display, hostname in device_mapping:
        row_cells = table.add_row().cells
        row_cells[0].text = username_display
        row_cells[1].text = hostname

    filename = f'{datetime.now().strftime("%Y%m%d")}_{filename_prefix}_devices.docx'
    file_path = os.path.join(settings.MEDIA_ROOT, 'exports', filename)

    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)

    return file_path

class CreateContest(PermissionRequiredMixin, TitleMixin, CreateView):
    template_name = 'contest/create.html'
    model = Contest
    form_class = ContestForm
    permission_required = 'judge.add_contest'
    permission_denied_message = _('You are not allowed to create contests.')

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        kwargs['org_slug'] = self.kwargs.get('slug')
        return kwargs

    def get_title(self):
        return _('Create new contest')

    def get_content_title(self):
        return _('Create new contest')

    def get_contest_problem_formset(self):
        if self.request.POST:
            return ProposeContestProblemFormSet(self.request.POST, form_kwargs={'user': self.request.user})
        return ProposeContestProblemFormSet(form_kwargs={'user': self.request.user})

    def get_context_data(self, **kwargs):
        data = super().get_context_data(**kwargs)
        data['contest_problem_formset'] = self.get_contest_problem_formset()
        if hasattr(self, "organization") and self.organization:
            data["users_json"] = json.dumps([
                {
                    "id": m.id,
                    "name": f"{m.user.username} ({m.user.first_name})"
                }
                for m in self.organization.members.all()
                    .filter(is_unlisted=False)
                    .exclude(user__is_staff=True)
                    .select_related("user")
            ])
        else:
            data["users_json"] = json.dumps([])

        return data

    def save_contest_form(self, form):
        self.object = form.save(commit=False)   # BẮT BUỘC

        self.object.save()
        form.save_m2m()  
        rooms = form.cleaned_data.get("exam_rooms")
        if rooms:
            self.object.exam_room.set(rooms)

        self.object.authors.add(self.request.profile)

    def post(self, request, *args, **kwargs):
        self.object = None
        post_data = request.POST.copy()
        is_exam = post_data.pop('is_exam', False)
        create_accounts = post_data.pop('create_exam_accounts', False)
        form = ContestForm(post_data, instance=self.object, org_slug=self.kwargs.get('slug'))
        form_set = self.get_contest_problem_formset()

        if form.is_valid() and form_set.is_valid():
            with revisions.create_revision(atomic=True):
                self.save_contest_form(form)
                for problem in form_set.save(commit=False):
                    problem.contest = self.object
                    problem.save()
                revisions.set_comment(_('Created on site'))
                revisions.set_user(self.request.user)
            on_new_contest.delay(self.object.key)

            if is_exam:
                tag, created = ContestTag.objects.get_or_create(
                    name="exam",
                    defaults={'color': '#ea2f2e'}
                )
                self.object.tags.add(tag)


            if create_accounts and is_exam:
                try:
                    organization = Organization.objects.get(slug=self.kwargs['slug'])       
                    self.object.organizations.add(organization)
                    problems = list(self.object.problems.all())
                    members = list(organization.members.filter(is_unlisted=False).exclude(user__username__regex=r'^[A-Z]{4}_').select_related('user'))
                    base_usernames = [m.user.username for m in members if not m.user.is_staff]
                    random_prefixes = [random_upper(4) for _ in base_usernames]
                    new_usernames = [f"{prefix}_{username}" for prefix, username in zip(random_prefixes, base_usernames)]
                    existing_usernames = set(get_user_model().objects.filter(username__in=new_usernames).values_list('username', flat=True))
                    username_map = {}

                    for base_username in base_usernames:
                        new_username = f"{random_upper(4)}_{base_username}"
                        while new_username in existing_usernames or new_username in username_map.values():
                            new_username = f"{random_upper(4)}_{base_username}"
                        username_map[base_username] = new_username
                        existing_usernames.add(new_username)

                    new_users = []
                    profiles = []
                    new_accounts = []
                    
                    SECRET_KEY = settings.SECRET_KEY
                    def get_fernet_key(secret):
                        key = hashlib.sha256(secret.encode()).digest()
                        return base64.urlsafe_b64encode(key)
                        
                    fernet = Fernet(get_fernet_key(SECRET_KEY))
                    with open("config/prehashed_p.encrypted", "rb") as f:
                        encrypted_data = f.read()
                    decrypted = fernet.decrypt(encrypted_data).decode()
                    password_pool = json.loads(decrypted)

                    for member in members:
                        old_user = member.user
                        if not old_user.is_staff:
                            new_username = username_map[old_user.username]
                            selected = random.choice(password_pool)
                            raw_password = selected["raw"]
                            hashed_password = selected["hashed"]
                            new_user = get_user_model()(
                                username=new_username,
                                password=hashed_password,
                            )
                            new_users.append(new_user)
                            profiles.append(Profile(user=new_user, timezone=member.timezone, language=member.language))
                            new_accounts.append((new_username, raw_password))
                    
                    get_user_model().objects.bulk_create(new_users)
                    Profile.objects.bulk_create(profiles)
                  
                    for profile in profiles:
                        profile.organizations.add(organization)

                    self.object.private_contestants.add(*profiles)
                    link_download = export_accounts_to_docx(new_accounts, f'{organization.slug}_{self.object}')
                    request.session['account_docx_path'] = link_download
                except Organization.DoesNotExist:
                    print("[ERROR] Không tìm thấy tổ chức.")
                    organization = None

            rooms = form.cleaned_data.get("exam_rooms")

            if rooms and is_exam:
                room_user_map_raw = request.POST.get("room_user_map")
                room_user_map = {}

                if room_user_map_raw:
                    try:
                        room_user_map = json.loads(room_user_map_raw)
                    except Exception:
                        room_user_map = {}

                contest_seats = []
                export_mapping = []

                # flatten user đã được map
                mapped_user_ids = set()
                for room_id, user_ids in room_user_map.items():
                    mapped_user_ids.update(map(int, user_ids))

                from collections import defaultdict
                devices_by_room = defaultdict(list)
                from django.utils import timezone as tz
                now_start = self.object.start_time
                now_end = self.object.end_time

                busy_device_ids = ContestSeat.objects.filter(
                    contest__start_time__lt=now_end,
                    contest__end_time__gt=now_start,
                    device__isnull=False 
                ).exclude(
                    contest=self.object
                ).values_list('device_id', flat=True)

                available_devices = list(
                    Device.objects.filter(
                        room__in=rooms,
                        is_active=True
                    ).exclude(
                        id__in=busy_device_ids
                    ).order_by("room__code", "hostname")
                )
                for d in available_devices:
                    devices_by_room[d.room_id].append(d)

                organization = Organization.objects.get(slug=self.kwargs['slug'])
                self.object.organizations.add(organization)

                if self.object.is_private:
                    members = list(
                        self.object.private_contestants
                        .filter(is_unlisted=False)
                        .exclude(user__is_staff=True)
                        .exclude(user=self.request.user)
                        .select_related('user')
                    )
                else:
                    members = list(
                        organization.members
                        .filter(is_unlisted=False)
                        .exclude(user__is_staff=True)
                        .exclude(user=self.request.user)
                        .select_related('user')
                    )

                if room_user_map:
                    for room_id, user_ids in room_user_map.items():
                        for uid in user_ids:
                            found = any(m.id == int(uid) for m in members)
                        room_id = int(room_id)
                        users_in_room = [m for m in members if m.id in map(int, user_ids)]

                        devices = devices_by_room.get(room_id, [])

                        if not devices:
                            continue

                        step = max(1, len(devices) // max(1, len(users_in_room)))
                        selected = devices[::step][:len(users_in_room)]

                        for member, device in zip(users_in_room, selected):
                            contest_seats.append(ContestSeat(
                                contest=self.object,
                                user=member,
                                device=device
                            ))

                            export_mapping.append((
                                f"{member.user.username} ({member.user.first_name})",
                                device.hostname
                            ))

                    ContestSeat.objects.bulk_create(contest_seats)

                else:        
                    try:
                        if len(available_devices) < len(members):
                            raise Exception(
                                f"Not enough available devices: {len(available_devices)} available, {len(members)} members"
                            )

                        step = max(1, len(available_devices) // len(members))
                        selected_devices = available_devices[::step][:len(members)]
                        contest_seats = []
                        export_mapping = []

                        for member, device in zip(members, selected_devices):
                            contest_seats.append(ContestSeat(
                                contest=self.object,
                                user=member,
                                device=device
                            ))
                            export_mapping.append((
                                f"{member.user.username} ({member.user.first_name})",
                                device.hostname
                            ))

                        ContestSeat.objects.bulk_create(contest_seats)
                    except Organization.DoesNotExist:
                        print("[ERROR] Không tìm thấy tổ chức.")
 
            return HttpResponseRedirect(self.get_success_url())
        else:
            return self.render_to_response(self.get_context_data(*args, **kwargs))

    def dispatch(self, request, *args, **kwargs):
        try:
            return super().dispatch(request, *args, **kwargs)
        except PermissionDenied as e:
            return generic_message(request, _('Permission denied'), e)


class EditContest(ContestMixin, LoginRequiredMixin, TitleMixin, UpdateView):
    template_name = 'contest/edit.html'
    model = Contest
    form_class = ContestForm

    def get_object(self, queryset=None):
        contest = super(EditContest, self).get_object(queryset)
        if not contest.is_editable_by(self.request.user):
            raise PermissionDenied(_('You are not allowed to edit this contest.'))
        return contest

    def get_form_kwargs(self):
        kwargs = super(EditContest, self).get_form_kwargs()
        # Due to some limitation with query set in select2
        # We only support this if the contest is private for only
        # 1 organization
        if self.object.organizations.count() == 1:
            kwargs['org_pk'] = self.object.organizations.values_list('pk', flat=True)[0]

        kwargs['user'] = self.request.user
        return kwargs

    def get_title(self):
        return _('Editing contest {0}').format(self.object.name)

    def get_content_title(self):
        return mark_safe(escape(_('Editing contest %s')) % (
            format_html('<a href="{1}">{0}</a>', self.object.name,
                        reverse('contest_view', args=[self.object.key]))))

    def get_contest_problem_formset(self):
        if self.request.POST:
            return ProposeContestProblemFormSet(self.request.POST, instance=self.get_object(),
                                                form_kwargs={'user': self.request.user})
        return ProposeContestProblemFormSet(instance=self.get_object(), form_kwargs={'user': self.request.user})

    def get_context_data(self, **kwargs):
        data = super().get_context_data(**kwargs)
        data['contest_problem_formset'] = self.get_contest_problem_formset()
        return data

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        form = self.get_form()
        form_set = self.get_contest_problem_formset()

        if form.is_valid() and form_set.is_valid():
            with revisions.create_revision(atomic=True):
                form.save()
                problems = form_set.save(commit=False)

                for problem in form_set.deleted_objects:
                    problem.delete()

                for problem in problems:
                    problem.contest = self.object
                    problem.save()

                revisions.set_comment(_('Edited from site'))
                revisions.set_user(self.request.user)

            return HttpResponseRedirect(self.get_success_url())
        else:
            return self.render_to_response(self.get_context_data(object=self.object))


class ContestDataMixin(ContestMixin, LoginRequiredMixin):
    @cached_property
    def data_path(self):
        return os.path.join(settings.DMOJ_CONTEST_DATA_CACHE, '%s.zip' % self.object.id)

    def get_object(self, queryset=None):
        if not settings.DMOJ_CONTEST_DATA_DOWNLOAD:
            raise Http404()
        contest = super().get_object(queryset)
        if not contest.is_editable_by(self.request.user):
            raise PermissionDenied(_('You are not allowed to edit this contest.'))
        if not contest.ended:
            raise PermissionDenied(_('Please wait until the contest has ended to download data.'))
        return contest


class ContestPrepareData(ContestDataMixin, TitleMixin, SingleObjectMixin, FormView):
    title = gettext_lazy('Download contest data')
    template_name = 'contest/prepare-data.html'
    form_class = ContestDownloadDataForm

    @cached_property
    def _now(self):
        return timezone.now()

    @cached_property
    def can_prepare_data(self):
        return (
            self.object.data_last_downloaded is None or
            self.object.data_last_downloaded + settings.DMOJ_CONTEST_DATA_DOWNLOAD_RATELIMIT < self._now or
            not os.path.exists(self.data_path)
        )

    @cached_property
    def data_cache_key(self):
        return 'celery_status_id:contest_data_download_%s' % self.object.id

    @cached_property
    def in_progress_url(self):
        status_id = cache.get(self.data_cache_key)
        status = task_status_by_id(status_id).status if status_id else None
        return (
            self.build_task_url(status_id)
            if status in ('PENDING', 'PROGRESS', 'STARTED')
            else None
        )

    def build_task_url(self, status_id):
        return task_status_url_by_id(
            status_id,
            message=_('Preparing data for %s...') % (self.object.name,),
            redirect=reverse('contest_prepare_data', args=(self.object.key,)),
        )

    def form_valid(self, form):
        self.object.data_last_downloaded = self._now
        self.object.save()
        status = prepare_contest_data.delay(self.object.id, json.dumps(form.cleaned_data))
        cache.set(self.data_cache_key, status.id)
        return HttpResponseRedirect(self.build_task_url(status.id))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['can_prepare_data'] = self.can_prepare_data
        context['can_download_data'] = os.path.exists(self.data_path)
        context['in_progress_url'] = self.in_progress_url
        context['ratelimit'] = settings.DMOJ_CONTEST_DATA_DOWNLOAD_RATELIMIT

        if not self.can_prepare_data:
            context['time_until_can_prepare'] = (
                settings.DMOJ_CONTEST_DATA_DOWNLOAD_RATELIMIT - (self._now - self.object.data_last_downloaded)
            )
        return context

    def get(self, request, *args, **kwargs):
        self.object = self.get_object()
        return super().get(request, *args, **kwargs)

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        if not self.can_prepare_data or self.in_progress_url is not None:
            raise PermissionDenied('You are not allowed to prepare new data.')
        return super().post(request, *args, **kwargs)


class ContestDownloadData(ContestDataMixin, SingleObjectMixin, View):
    def get(self, request, *args, **kwargs):
        self.object = self.get_object()

        if not os.path.exists(self.data_path):
            raise Http404()

        response = HttpResponse()

        if hasattr(settings, 'DMOJ_CONTEST_DATA_INTERNAL'):
            url_path = '%s/%s.zip' % (settings.DMOJ_CONTEST_DATA_INTERNAL, self.object.id)
        else:
            url_path = None
        add_file_response(request, response, url_path, self.data_path)

        response['Content-Type'] = 'application/zip'
        response['Content-Disposition'] = 'attachment; filename=%s-data.zip' % self.object.key
        return response

class ContestEndExam(LoginRequiredMixin, ContestMixin, SingleObjectMixin, View):
    def dispatch(self, request, *args, **kwargs):
        if request.method != 'POST':
            return HttpResponseForbidden()

        return super(ContestEndExam, self).dispatch(request, *args, **kwargs)

    def post(self, request, *args, **kwargs):
        contest = self.get_object()
        # if request.is_seb and not (request.user.is_staff or request.user.is_superuser):
        profile = request.profile
        if profile.current_contest is None or profile.current_contest.contest_id != contest.id:
            return generic_message(request, _('No such contest'),
                                _('You are not in contest "%s".') % contest.key, 404)
        participation = profile.current_contest

        participation.finished_at = timezone.now()
        participation.save()
        return HttpResponseRedirect('/quit_seb')

class ContestDevices(ContestMixin, TitleMixin, TemplateView):
    template_name = 'contest/devices.html'

    def get_title(self):
        return _('%s Devices') % self.object.name

    @cached_property
    def can_edit(self):
        return self.object.is_editable_by(self.request.user)

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()

        if not self.can_edit:
            raise PermissionDenied()

        return super().dispatch(request, *args, **kwargs)

    def get_object(self):
        return get_object_or_404(Contest, key=self.kwargs['contest'])
        
    @cached_property
    def seats(self):
        return (
            ContestSeat.objects
            .filter(contest=self.object)
            .select_related('user__user', 'device__room')
            .order_by('device__hostname')
        )

    @cached_property
    def devices(self):
        return Device.objects.filter(
            room__in=self.object.exam_room.all(),
            is_active=True
        ).select_related('room').order_by('hostname')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.request.user.is_authenticated:
            try:
                context['live_participation'] = (
                    self.request.profile.contest_history.get(
                        contest=self.object,
                        virtual=ContestParticipation.LIVE,
                    )
                )
            except ContestParticipation.DoesNotExist:
                context['live_participation'] = None
                context['has_joined'] = False
            else:
                context['has_joined'] = True
        else:
            context['live_participation'] = None
            context['has_joined'] = False

        context['contest'] = self.object
        context['seats'] = self.seats
        context['devices'] = self.devices
        context['can_edit'] = self.can_edit
        context['now'] = self.object._now
        context["devices_map"] = {}

        busy_device_ids = set(
            ContestSeat.objects
            .filter(
                contest__start_time__lt=self.object.end_time,
                contest__end_time__gt=self.object.start_time,
                device__isnull=False,
            )
            .exclude(contest=self.object)
            .exclude(contest__end_time=self.object.start_time)
            .exclude(contest__start_time=self.object.end_time)
            .values_list("device_id", flat=True)
        )
        current_contest_seat_device_ids = {
            seat.device_id for seat in self.seats if seat.device_id
        }

        for seat in self.seats:
            occupied_by_others = (
                busy_device_ids | (current_contest_seat_device_ids - {seat.device_id})
            )
            context["devices_map"][seat.id] = self.devices.exclude(
                id__in=occupied_by_others
            )

        return context

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()

        seat_id = request.POST.get("seat_id")
        device_id = request.POST.get("device_id")

        try:
            seat = ContestSeat.objects.get(id=seat_id, contest=self.object)

            if device_id:
                # validate device
                if not Device.objects.filter(
                    id=device_id,
                    room__in=self.object.exam_room.all(),
                    is_active=True
                ).exists():
                    raise Exception("Invalid device")

                exists = ContestSeat.objects.filter(
                    contest=self.object,
                    device_id=device_id
                ).exclude(id=seat.id).exists()

                if exists:
                    raise Exception("This device is already assigned to another user.")
                
                conflict = ContestSeat.objects.filter(
                    device_id=device_id,
                    contest__start_time__lt=self.object.end_time,
                    contest__end_time__gt=self.object.start_time,
                ).exclude(contest=self.object).exists()

                if conflict:
                    raise Exception("Device is already used in another overlapping contest")

                seat.device_id = device_id
                seat.updated_at = timezone.now()
            else:
                seat.device = None
                seat.updated_at = timezone.now()

            seat.save()

        except Exception as e:
            # print("[DEVICE UPDATE ERROR]", e)
            messages.error(request, str(e)) 

        return redirect('contest_device', contest=self.object.key)
